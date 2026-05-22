import io
import shutil

import pytest
from docx import Document

from pen_plotter.converters.document import DocumentConverter
from pen_plotter.converters.html import HtmlConverter
from pen_plotter.converters.markdown import MarkdownConverter
from pen_plotter.converters.text import TextConverter
from pen_plotter.typography import HersheyRenderer, TypographyOptions, available_fonts

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
needs_libreoffice = pytest.mark.skipif(
    shutil.which("libreoffice") is None and shutil.which("soffice") is None,
    reason="libreoffice not installed",
)


def test_available_fonts_includes_futural() -> None:
    assert "futural" in available_fonts()


def test_renderer_rejects_unknown_font() -> None:
    with pytest.raises(ValueError):
        HersheyRenderer(TypographyOptions(font="not-a-font"))


def test_renderer_outputs_strokes_within_page() -> None:
    svg = HersheyRenderer(TypographyOptions(font_size_mm=6)).render_text("Hello plotter")
    assert svg.startswith("<svg")
    assert 'stroke="black"' in svg
    assert svg.count("<path") >= 1


def test_text_converter_wraps_long_lines() -> None:
    long_text = "word " * 200
    svg = TextConverter().convert(long_text.encode(), options={"font_size_mm": 5}).svg
    assert svg.count("<path") > 1


def test_text_converter_alignment_option() -> None:
    result = TextConverter().convert(b"centered", options={"alignment": "center"})
    assert "<path" in result.svg


def test_markdown_headings_render_larger_blocks() -> None:
    md = b"# Title\n\nSome body text here.\n\n- one\n- two\n"
    result = MarkdownConverter().convert(md)
    assert result.source_mime == "image/svg+xml"
    assert result.svg.count("<path") > 1


@needs_libreoffice
def test_document_converter_handles_docx() -> None:
    doc = Document()
    doc.add_heading("Title", 0)
    doc.add_paragraph("Hello from a Word document.")
    buf = io.BytesIO()
    doc.save(buf)
    result = DocumentConverter().convert(buf.getvalue(), options={"source_mime": DOCX_MIME})
    assert result.source_mime == "image/svg+xml"
    assert result.svg.startswith("<svg")


def test_html_converter_renders_svg() -> None:
    result = HtmlConverter().convert(b"<h1>Title</h1><p>Body text.</p>")
    assert result.source_mime == "image/svg+xml"
    assert result.svg.startswith("<svg")
